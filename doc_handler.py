# -*- coding: utf-8 -*-
"""
文档处理器
负责读取和写入 .docx 文件，保留原始格式。
"""
import os
import shutil
from dataclasses import dataclass
from typing import Optional


@dataclass
class ParagraphInfo:
    """段落信息"""
    index: int
    text: str
    style_name: str = ""
    is_heading: bool = False
    heading_level: int = 0
    word_count: int = 0


def read_docx(path: str) -> tuple:
    """
    读取 docx 文件。
    返回 (Document对象, 段落信息列表)。
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = Document(path)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        is_heading = style.lower().startswith('heading') or '标题' in style
        level = 0
        if is_heading:
            for ch in style:
                if ch.isdigit():
                    level = int(ch)
                    break

        paragraphs.append(ParagraphInfo(
            index=i,
            text=text,
            style_name=style,
            is_heading=is_heading,
            heading_level=level,
            word_count=len(text),
        ))

    return doc, paragraphs


def replace_paragraph_text(doc, index: int, new_text: str) -> bool:
    """
    替换指定段落的文本内容，保留第一个run的格式。
    """
    if index >= len(doc.paragraphs):
        return False

    para = doc.paragraphs[index]
    if not para.text.strip():
        return False

    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    return True


def save_docx(doc, output_path: str, also_copy_as: Optional[str] = None):
    """保存文档，可选同时复制一份带中文名的版本。"""
    doc.save(output_path)
    if also_copy_as:
        shutil.copy2(output_path, also_copy_as)


def get_content_paragraphs(paragraphs: list[ParagraphInfo]) -> list[ParagraphInfo]:
    """过滤出正文段落（排除标题和空段落）"""
    return [p for p in paragraphs if not p.is_heading and p.word_count > 0]


def analyze_document(paragraphs: list[ParagraphInfo]) -> dict:
    """分析文档基本统计信息"""
    total = len(paragraphs)
    content = get_content_paragraphs(paragraphs)
    headings = [p for p in paragraphs if p.is_heading]
    empty = [p for p in paragraphs if p.word_count == 0]
    total_words = sum(p.word_count for p in content)

    return {
        'total_paragraphs': total,
        'content_paragraphs': len(content),
        'headings': len(headings),
        'empty_paragraphs': len(empty),
        'total_words': total_words,
        'avg_words_per_para': total_words // max(len(content), 1),
    }
